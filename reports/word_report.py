"""
Generador del informe Word del Estudio de Mercado — Kioskos Parque Bicentenario.

Recibe un dict `charts` con las imágenes ya renderizadas como bytes PNG.
Esto evita llamar a kaleido dentro del contexto de Streamlit.
"""

import io
import os
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── helpers ───────────────────────────────────────────────────────────────────

def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _p(doc, text, size=11, bold=False):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return para


def _insert_img(doc, png_bytes, caption, width=6.0):
    """Inserta imagen PNG (bytes) con pie de figura."""
    if not png_bytes:
        return
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].italic    = True
    doc.add_paragraph()


def _kpi_table(doc, rows):
    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.style = "Table Grid"
    for j, h in enumerate(["Indicador", "Valor"]):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (lbl, val) in enumerate(rows, 1):
        tbl.rows[i].cells[0].text = lbl
        tbl.rows[i].cells[1].text = val
        for j in range(2):
            tbl.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()


def _bullet(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


# ── chart generator con matplotlib (sin kaleido, funciona en Streamlit) ───────

def build_charts(enc, kpis, stats, fc):
    """
    Genera todas las imágenes con matplotlib puro (sin kaleido/subprocesos).
    Retorna dict: nombre → bytes PNG.
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.ticker as mticker
    import numpy as np
    from models.kiosko_model import PLAN_FASES, COMMERCIAL_PARAMS

    BLUE   = "#1a3a5c"
    ORANGE = "#e67e22"
    GREEN  = "#27ae60"
    COLS   = [BLUE, ORANGE, GREEN, "#8e44ad", "#e74c3c", "#16a085", "#f39c12"]

    def _save(fig):
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    facecolor="white", edgecolor="none")
        buf.seek(0)
        data = buf.read()
        plt.close(fig)
        return data

    def _pie(series, title, w=5, h=4):
        vc = series.dropna().str.strip().value_counts()
        fig, ax = plt.subplots(figsize=(w, h))
        wedges, texts, autotexts = ax.pie(
            vc.values, labels=vc.index, autopct="%1.1f%%",
            colors=COLS[:len(vc)], startangle=90,
            wedgeprops=dict(linewidth=1, edgecolor="white"),
        )
        for t in autotexts:
            t.set_fontsize(9)
        ax.set_title(title, fontsize=12, fontweight="bold", color=BLUE, pad=12)
        return _save(fig)

    def _hbar(labels, values, title, w=6, h=4, color=None):
        """Barra horizontal en porcentaje con etiquetas de datos."""
        color = color or BLUE
        total = sum(values)
        pcts  = [v / total * 100 for v in values] if total > 0 else list(values)
        fig, ax = plt.subplots(figsize=(w, h))
        bars = ax.barh(labels, pcts, color=color, alpha=0.85, height=0.55)
        ax.set_xlabel("% de encuestados", fontsize=9)
        ax.set_title(title, fontsize=12, fontweight="bold", color=BLUE, pad=10)
        ax.spines[["top","right"]].set_visible(False)
        for bar, pct in zip(bars, pcts):
            ax.text(bar.get_width() + max(pcts) * 0.01,
                    bar.get_y() + bar.get_height() / 2,
                    f"{pct:.1f}%", va="center", fontsize=9, fontweight="bold")
        ax.set_xlim(0, max(pcts) * 1.22)
        ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:.0f}%"))
        fig.tight_layout()
        return _save(fig)

    charts = {}

    # ── 1. Género ──────────────────────────────────────────────────────────
    try:
        if "genero" in enc.columns:
            charts["genero"] = _pie(enc["genero"], "Distribución por género")
    except Exception:
        charts["genero"] = None

    # ── 2. Edad ───────────────────────────────────────────────────────────
    try:
        if "edad" in enc.columns:
            orden = ["< 20", "20–29", "30–39", "40–49", "50–59", "60+"]
            bins  = pd.cut(enc["edad"].dropna(),
                           bins=[0,20,30,40,50,60,120], labels=orden)
            vc = bins.value_counts().reindex(orden).fillna(0)
            charts["edad"] = _hbar(vc.index.tolist(), vc.values.tolist(),
                                   "Distribución por grupo de edad")
    except Exception:
        charts["edad"] = None

    # ── 3. Motivo de visita ───────────────────────────────────────────────
    try:
        if "motivo_visita" in enc.columns:
            charts["motivo"] = _pie(enc["motivo_visita"], "Motivo principal de visita")
    except Exception:
        charts["motivo"] = None

    # ── 4. Tráfico mensual 2025 ───────────────────────────────────────────
    try:
        meses_dfs = [
            df[["mes", "num_mes", "total"]].copy()
            for df in stats.values()
            if "total" in df.columns and "mes" in df.columns
        ]
        agg = (pd.concat(meses_dfs)
               .groupby(["mes", "num_mes"])["total"].sum()
               .reset_index().sort_values("num_mes"))
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(agg["mes"], agg["total"], color=BLUE, linewidth=2.2,
                marker="o", markersize=5)
        ax.fill_between(range(len(agg)), agg["total"].values,
                        alpha=0.12, color=BLUE)
        media = agg["total"].mean()
        ax.axhline(media, color="#e74c3c", linewidth=1.4, linestyle="--",
                   label=f"Media: {int(media):,}")
        ax.legend(fontsize=9)
        ax.set_xticks(range(len(agg)))
        ax.set_xticklabels(agg["mes"], rotation=35, ha="right", fontsize=8)
        ax.set_xlabel("Mes", fontsize=9)
        ax.set_ylabel("Visitas", fontsize=9)
        ax.set_title("Tráfico de visitantes al Parque Bicentenario — 2025",
                     fontsize=12, fontweight="bold", color=BLUE)
        ax.yaxis.set_major_formatter(mticker.FuncFormatter(
            lambda x, _: f"{int(x/1000)}k" if x >= 1000 else str(int(x))))
        ax.spines[["top","right"]].set_visible(False)
        fig.tight_layout()
        charts["trafico"] = _save(fig)
    except Exception:
        charts["trafico"] = None

    # ── 5. Demanda vs plan de kioskos ─────────────────────────────────────
    # Metodología: se ancla en las visitas reales medidas en 2025 (vis_dia) y se
    # proyecta con el factor de crecimiento poblacional del sector (mismo método
    # que fig_ingresos_fases_zona en models/kiosko_model.py), NO con un % fijo
    # arbitrario de la población total del sector.
    try:
        cons_pct = kpis["consumiria_pct"] / 100
        vis_dia  = round(kpis["total_visitas_parque_2025"] / 365)
        fases = sorted(PLAN_FASES.items())
        anios = [str(a) for a, _ in fases]
        kios  = [sum(p["kioskos"].values()) for _, p in fases]
        pob_base = fases[0][1]["poblacion_hab"]
        factores = [p["poblacion_hab"] / pob_base for _, p in fases]
        demanda  = [round(vis_dia * f * cons_pct) for f in factores]

        x = np.arange(len(anios))
        w = 0.38
        fig, ax = plt.subplots(figsize=(7, 4))
        capacidad = [k * 30 for k in kios]
        b1 = ax.bar(x - w/2, demanda, w, label="Consumidores potenciales/día",
                    color=BLUE, alpha=0.85)
        b2 = ax.bar(x + w/2, capacidad, w,
                    label="Capacidad diaria (kioskos × 30 tx)", color=ORANGE, alpha=0.85)
        # etiquetas de datos
        _top = max(max(demanda), max(capacidad)) * 0.015
        for bar, val in zip(b1, demanda):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + _top,
                    f"{val:,}", ha="center", va="bottom", fontsize=8, fontweight="bold", color=BLUE)
        for bar, val in zip(b2, capacidad):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + _top,
                    f"{val:,}", ha="center", va="bottom", fontsize=8, fontweight="bold", color=ORANGE)
        ax.set_xticks(x)
        ax.set_xticklabels([f"Fase\n{a}" for a in anios], fontsize=9)
        ax.set_ylabel("Personas / transacciones diarias", fontsize=9)
        ax.set_ylim(0, max(max(demanda), max(capacidad)) * 1.22)
        ax.set_title("Demanda potencial vs. capacidad del plan de kioskos",
                     fontsize=11, fontweight="bold", color=BLUE)
        ax.legend(fontsize=9)
        ax.spines[["top","right"]].set_visible(False)
        fig.tight_layout()
        charts["demanda"] = _save(fig)
    except Exception:
        charts["demanda"] = None

    # ── 6. Visitantes vs consumidores por fase ────────────────────────────
    # Misma metodología: vis_dia real 2025 proyectado con el factor de
    # crecimiento poblacional del sector por fase.
    try:
        cons_pct2 = kpis["consumiria_pct"] / 100
        vis_dia2  = round(kpis["total_visitas_parque_2025"] / 365)
        fases = sorted(PLAN_FASES.items())
        anios = [str(a) for a, _ in fases]
        pob_base2 = fases[0][1]["poblacion_hab"]
        vis   = [round(vis_dia2 * (p["poblacion_hab"] / pob_base2)) for _, p in fases]
        cons  = [round(v * cons_pct2) for v in vis]
        x = np.arange(len(anios))
        w = 0.38
        fig, ax = plt.subplots(figsize=(7, 4))
        bv = ax.bar(x - w/2, vis,  w, label="Visitantes/día estimados", color=BLUE, alpha=0.85)
        bc = ax.bar(x + w/2, cons, w, label=f"Consumidores ({kpis['consumiria_pct']}%)", color=GREEN, alpha=0.85)
        _top2 = max(max(vis), max(cons)) * 0.015
        for bar, val in zip(bv, vis):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + _top2,
                    f"{val:,}", ha="center", va="bottom", fontsize=8, fontweight="bold", color=BLUE)
        for bar, val in zip(bc, cons):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + _top2,
                    f"{val:,}", ha="center", va="bottom", fontsize=8, fontweight="bold", color=GREEN)
        ax.set_xticks(x)
        ax.set_xticklabels([f"Fase\n{a}" for a in anios], fontsize=9)
        ax.set_ylabel("Personas por día", fontsize=9)
        ax.set_ylim(0, max(max(vis), max(cons)) * 1.22)
        ax.set_title("Proyección de visitantes vs. consumidores por fase",
                     fontsize=11, fontweight="bold", color=BLUE)
        ax.legend(fontsize=9)
        ax.spines[["top","right"]].set_visible(False)
        fig.tight_layout()
        charts["visitantes_vs"] = _save(fig)
    except Exception:
        charts["visitantes_vs"] = None

    # ── 7. Kioskos por fase ───────────────────────────────────────────────
    try:
        fases = sorted(PLAN_FASES.items())
        anios = [str(a) for a, _ in fases]
        kios  = [sum(p["kioskos"].values()) for _, p in fases]
        fig, ax = plt.subplots(figsize=(7, 4))
        bars = ax.bar(anios, kios, color=BLUE, alpha=0.85, width=0.5)
        for bar, val in zip(bars, kios):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                    str(val), ha="center", va="bottom", fontweight="bold", fontsize=11)
        ax.set_ylabel("Número de kioskos", fontsize=9)
        ax.set_title("Proyección de kioskos — Plan por fases 2026–2036",
                     fontsize=11, fontweight="bold", color=BLUE)
        ax.set_ylim(0, max(kios) * 1.2)
        ax.spines[["top","right"]].set_visible(False)
        fig.tight_layout()
        charts["kioskos"] = _save(fig)
    except Exception:
        charts["kioskos"] = None

    # ── 8. Ingresos por zona ──────────────────────────────────────────────
    # Ancla en visitas reales 2025 + factor de crecimiento poblacional (misma
    # metodología que fig_ingresos_fases_zona en models/kiosko_model.py).
    try:
        gasto = kpis["gasto_promedio_usd"]
        cons_pct3 = kpis["consumiria_pct"] / 100
        vis_dia3  = round(kpis["total_visitas_parque_2025"] / 365)
        fases = sorted(PLAN_FASES.items())
        anios = [str(a) for a, _ in fases]
        pob_base3 = fases[0][1]["poblacion_hab"]
        ingresos_zona = []
        for a, p in fases:
            factor = p["poblacion_hab"] / pob_base3
            vis_d  = round(vis_dia3 * factor)
            cons_d = round(vis_d * cons_pct3)
            ing    = round(cons_d * gasto * 365 / 1000)
            ingresos_zona.append(ing)
        fig, ax = plt.subplots(figsize=(7, 4))
        bars = ax.bar(anios, ingresos_zona, color=ORANGE, alpha=0.85, width=0.5)
        for bar, val in zip(bars, ingresos_zona):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(ingresos_zona)*0.01,
                    f"${val:,}k", ha="center", va="bottom", fontsize=9, fontweight="bold")
        ax.set_ylabel("Ingresos potenciales (miles USD/año)", fontsize=9)
        ax.set_title("Proyección de ingresos totales de la zona — 2026–2036",
                     fontsize=11, fontweight="bold", color=BLUE)
        ax.spines[["top","right"]].set_visible(False)
        ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"${int(x):,}k"))
        fig.tight_layout()
        charts["ing_zona"] = _save(fig)
    except Exception:
        charts["ing_zona"] = None

    # ── 9. Ingresos por kiosko ────────────────────────────────────────────
    # Misma metodología real (visitas 2025 × factor de crecimiento poblacional).
    try:
        gasto = kpis["gasto_promedio_usd"]
        cons_pct4 = kpis["consumiria_pct"] / 100
        vis_dia4  = round(kpis["total_visitas_parque_2025"] / 365)
        fases = sorted(PLAN_FASES.items())
        anios = [str(a) for a, _ in fases]
        pob_base4 = fases[0][1]["poblacion_hab"]
        ing_kio = []
        for a, p in fases:
            kios_n = sum(p["kioskos"].values())
            factor = p["poblacion_hab"] / pob_base4
            vis_d  = round(vis_dia4 * factor)
            cons_d = round(vis_d * cons_pct4)
            ing    = round(cons_d * gasto * 365 / kios_n / 1000, 1) if kios_n else 0
            ing_kio.append(ing)
        fig, ax = plt.subplots(figsize=(7, 4))
        bars = ax.bar(anios, ing_kio, color=GREEN, alpha=0.85, width=0.5)
        for bar, val in zip(bars, ing_kio):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(ing_kio)*0.01,
                    f"${val:,}k", ha="center", va="bottom", fontsize=9, fontweight="bold")
        ax.set_ylabel("Ingresos por kiosko (miles USD/año)", fontsize=9)
        ax.set_title("Proyección de ingresos promedio por kiosko — 2026–2036",
                     fontsize=11, fontweight="bold", color=BLUE)
        ax.spines[["top","right"]].set_visible(False)
        ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"${int(x):,}k"))
        fig.tight_layout()
        charts["ing_kio"] = _save(fig)
    except Exception:
        charts["ing_kio"] = None

    # ── 10. Distribución de kioskos en el bulevar ────────────────────────
    try:
        from models.kiosko_model import GIRO_ROTATION, GIROS
        rotation = GIRO_ROTATION["comercial_alta_densidad"]
        n = 10
        giros = rotation[:n]
        colores_giro = {
            "Bebidas":                            "#2196F3",
            "Comida rápida":                      "#FF5722",
            "Helados y café":                     "#9C27B0",
            "Snacks saludables":                  "#4CAF50",
            "Artículos para mascotas y souvenirs":"#FF9800",
            "Deportivo":                          "#F44336",
        }
        colors = [colores_giro.get(g, BLUE) for g in giros]

        fig, ax = plt.subplots(figsize=(10, 3))
        ax.set_xlim(0, n)
        ax.set_ylim(0, 1)
        ax.axis("off")

        for i, (giro, col) in enumerate(zip(giros, colors)):
            # Rectángulo del kiosko
            rect = plt.Rectangle((i + 0.1, 0.15), 0.8, 0.55,
                                  color=col, alpha=0.88, zorder=3)
            ax.add_patch(rect)
            # Número
            ax.text(i + 0.5, 0.42, f"K{i+1}",
                    ha="center", va="center", fontsize=9,
                    fontweight="bold", color="white", zorder=4)
            # Etiqueta giro (debajo)
            label = giro.replace(" y ", "\ny ").replace(" para ", "\npara ")
            ax.text(i + 0.5, 0.08, label,
                    ha="center", va="top", fontsize=6, color="#333", zorder=4)

        ax.set_title("Distribución comercial propuesta — Bulevar de las Canchas (10 kioskos, Fase 1)",
                     fontsize=10, fontweight="bold", color=BLUE, pad=6)
        ax.set_facecolor("#f0f0f0")
        fig.patch.set_facecolor("white")
        fig.tight_layout()
        charts["layout"] = _save(fig)
    except Exception:
        charts["layout"] = None

    # ── 11. Productos más demandados (enc["productos_interes"]) ──────────
    try:
        if "productos_interes" in enc.columns:
            vc_p = enc["productos_interes"].dropna().value_counts().head(8)
            charts["productos"] = _hbar(
                vc_p.index[::-1].tolist(),
                vc_p.values[::-1].tolist(),
                "Productos y servicios más demandados",
                w=7, h=4, color=ORANGE,
            )
        else:
            charts["productos"] = None
    except Exception:
        charts["productos"] = None

    # ── 12. Temas más solicitados (análisis de comentarios) ───────────────
    try:
        from analysis.sentiment import analyze_comments
        result_df = analyze_comments(enc)
        col_t = next((c for c in ["temas", "topic", "tema"] if c in result_df.columns), None)
        if result_df is not None and not result_df.empty and col_t:
            serie = result_df[col_t].dropna()
            if serie.apply(lambda x: isinstance(x, list)).any():
                serie = serie.explode()
            tc = serie[serie != "Otro"].value_counts().head(8)
            total_t = tc.sum()
            tc_pct  = (tc / total_t * 100) if total_t > 0 else tc
            fig, ax = plt.subplots(figsize=(7, 4))
            bars = ax.barh(tc_pct.index[::-1], tc_pct.values[::-1],
                           color=BLUE, alpha=0.85, height=0.55)
            ax.set_xlabel("% de menciones en comentarios", fontsize=9)
            ax.set_title("Temas más solicitados en comentarios de visitantes",
                         fontsize=11, fontweight="bold", color=BLUE)
            ax.spines[["top","right"]].set_visible(False)
            for bar, val in zip(bars, tc_pct.values[::-1]):
                ax.text(bar.get_width() + max(tc_pct.values) * 0.01,
                        bar.get_y() + bar.get_height() / 2,
                        f"{val:.1f}%", va="center", fontsize=9, fontweight="bold")
            ax.set_xlim(0, max(tc_pct.values) * 1.22)
            ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"{x:.0f}%"))
            fig.tight_layout()
            charts["topics"] = _save(fig)
        else:
            charts["topics"] = None
    except Exception:
        charts["topics"] = None

    return charts


# ── helpers de firma ──────────────────────────────────────────────────────────

def _firma_table(doc, roles):
    """Tabla de firmas: lista de (rol, nombre) — nombre en blanco = línea."""
    tbl = doc.add_table(rows=1, cols=len(roles))
    tbl.style = "Table Grid"
    row = tbl.rows[0]
    for i, (rol, nombre) in enumerate(roles):
        cell = row.cells[i]
        # Línea de firma
        p_line = cell.add_paragraph("_" * 28)
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Nombre/cargo
        p_nom = cell.add_paragraph(nombre if nombre else " ")
        p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nom.runs[0].font.size = Pt(9)
        # Rol
        p_rol = cell.add_paragraph(rol)
        p_rol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rol.runs[0].bold = True
        p_rol.runs[0].font.size = Pt(9)
    doc.add_paragraph()


# ── generador principal ────────────────────────────────────────────────────────

def generate_word_report(enc, kpis, stats, fc, charts=None):
    """
    Genera el informe Word completo con estructura institucional EPMMOP.

    Parámetros
    ----------
    enc    : DataFrame de encuesta limpia
    kpis   : dict de KPIs (de get_kpi_summary)
    stats  : dict de estadísticas de visitas
    fc     : resultado de forecast_kioskos()
    charts : dict nombre→bytes PNG (de build_charts).
    """
    from models.kiosko_model import PLAN_FASES, PROYECCION_PARAMS, GIROS, GIRO_ROTATION

    if charts is None:
        charts = build_charts(enc, kpis, stats, fc)

    # ── variables ────────────────────────────────────────────────────────────
    n_enc    = kpis["n_encuestados"]
    vis_2025 = kpis["total_visitas_parque_2025"]
    vis_dia  = round(vis_2025 / 365)
    cons_pct = kpis["consumiria_pct"]
    cons_dia = round(vis_dia * cons_pct / 100)
    gasto    = kpis["gasto_promedio_usd"]
    calif    = kpis["calificacion_promedio"]
    aprueba  = kpis["aprueba_kiosko_pct"]
    mejora   = kpis["mejora_experiencia_pct"]
    edad     = kpis["edad_promedio"]
    pob_2026 = PROYECCION_PARAMS["poblacion_sector_hab"][2026]
    pob_2036 = PROYECCION_PARAMS["poblacion_sector_hab"][2036]

    # ── documento ────────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # ━━━ PORTADA ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_heading("ESTUDIO DE MERCADO", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_heading("Kioskos Comerciales — Parque Bicentenario", 1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for txt, sz, bold in [
        ("Empresa Pública Metropolitana de Movilidad y Obras Públicas", 12, True),
        ("EPMMOP — Promoción de Servicios", 11, False),
        ("Gerencia de Desarrollo Urbano y Espacio Público", 10, False),
    ]:
        p = doc.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = bold
        p.runs[0].font.size = Pt(sz)

    doc.add_paragraph()
    _kpi_table(doc, [
        ("Código del documento",  "EPMMOP-PS-EM-2026-001"),
        ("Versión",               "1.0"),
        ("Fecha de emisión",      date.today().strftime("%d/%m/%Y")),
        ("Estado",                "Para revisión y aprobación"),
        ("Clasificación",         "Uso interno — EPMMOP"),
    ])
    doc.add_page_break()

    # ━━━ 1. INTRODUCCIÓN ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "1. INTRODUCCIÓN", 1)
    _p(doc, (
        "El presente documento constituye el Estudio de Mercado para la implementación de "
        "kioskos comerciales en el Parque Bicentenario de Quito, desarrollado por la Gerencia "
        "de Promoción de Servicios de la Empresa Pública Metropolitana de Movilidad y Obras "
        "Públicas (EPMMOP)."
    ))
    _p(doc, (
        "El estudio tiene como propósito determinar la viabilidad comercial y operativa de "
        "instalar módulos de venta de productos y servicios en el Bulevar de las Canchas, "
        "zona de mayor concentración peatonal del parque, con base en evidencia cuantitativa "
        "obtenida directamente de los visitantes y en el análisis de los registros históricos "
        "de afluencia del año 2025."
    ))
    _p(doc, (
        f"El estudio se fundamenta en una encuesta aplicada a {n_enc} visitantes del parque "
        f"y en el análisis de {vis_2025:,} registros de visita correspondientes al año 2025, "
        "proporcionados por la EPMMOP. Los resultados sirven de base técnica para la toma de "
        "decisiones sobre la implementación, el dimensionamiento y el modelo de gestión del "
        "proyecto de kioskos comerciales."
    ))
    doc.add_page_break()

    # ━━━ 2. ANTECEDENTES ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "2. ANTECEDENTES", 1)
    _p(doc, (
        "El Parque Bicentenario ocupa los terrenos del antiguo Aeropuerto Internacional "
        "Mariscal Sucre de Quito, cuya operación fue trasladada al nuevo aeropuerto de Tababela "
        "en febrero de 2013. Con aproximadamente 50 hectáreas de área verde, es el parque urbano "
        "más extenso de la ciudad y uno de los más significativos de la región andina."
    ))
    _p(doc, (
        "Administrado por la EPMMOP en virtud de la delegación otorgada por el Municipio del "
        "Distrito Metropolitano de Quito (MDMQ), el parque se ha consolidado como el principal "
        f"espacio de esparcimiento, deporte y cultura del norte de Quito, registrando "
        f"{vis_2025:,} visitas durante el año 2025."
    ))
    _p(doc, (
        "A pesar de su alta afluencia, el parque no cuenta con una oferta comercial formal que "
        "permita a los visitantes adquirir productos de alimentación, hidratación u otros "
        "servicios complementarios durante su permanencia. Esta situación genera una demanda "
        "insatisfecha que ha sido cuantificada en el presente estudio, y representa una "
        "oportunidad de desarrollo económico local alineada con los objetivos de la EPMMOP."
    ))
    _p(doc, (
        "En el marco del Plan de Gestión del Parque Bicentenario y de los proyectos de "
        "desarrollo urbano asociados —entre los que destaca la apertura de las estaciones de "
        "Metro Bicentenario y Andalucía proyectada para 2029— se identificó la necesidad de "
        "realizar un estudio técnico de mercado que sustente la decisión de implementar "
        "kioskos comerciales en el Bulevar de las Canchas."
    ))
    doc.add_page_break()

    # ━━━ 3. MARCO NORMATIVO O REFERENCIAL ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "3. MARCO NORMATIVO O REFERENCIAL", 1)
    _p(doc, (
        "La implementación de kioskos comerciales en el Parque Bicentenario se enmarca en el "
        "siguiente conjunto de instrumentos legales y normativos vigentes:"
    ))
    normas = [
        ("Constitución de la República del Ecuador (2008)",
         "Art. 264: Los gobiernos municipales tendrán entre sus competencias exclusivas la "
         "prestación de servicios públicos y la regulación del uso del suelo urbano."),
        ("Código Orgánico de Organización Territorial, Autonomía y Descentralización — COOTAD",
         "Regula las competencias de los gobiernos autónomos descentralizados, incluyendo la "
         "gestión y uso de espacios públicos urbanos."),
        ("Código Orgánico del Ambiente (COA)",
         "Establece los principios de sustentabilidad que deben observarse en intervenciones "
         "en espacios verdes y áreas de uso público."),
        ("Ordenanza Metropolitana N.° 0172 — Uso y Gestión del Suelo del DMQ",
         "Regula los usos compatibles con las zonas de parques y equipamientos urbanos, "
         "incluyendo actividades comerciales complementarias de baja intensidad."),
        ("Plan de Uso y Gestión del Suelo (PUGS) del DMQ — 2022",
         "Clasifica el Parque Bicentenario como área de equipamiento recreativo de escala "
         "metropolitana, permitiendo usos complementarios como servicios de alimentación "
         "y actividades comerciales en módulos móviles o semi-permanentes."),
        ("Acuerdo de creación de la EPMMOP — Ordenanza Metropolitana N.° 0165",
         "Define las competencias de la EPMMOP para la administración, gestión y "
         "aprovechamiento del Parque Bicentenario."),
        ("Norma Técnica de Espacio Público del DMQ",
         "Establece los parámetros de diseño, accesibilidad y ocupación para elementos "
         "instalados en espacios públicos, aplicables al dimensionamiento de kioskos."),
    ]
    tbl_n = doc.add_table(rows=len(normas) + 1, cols=2)
    tbl_n.style = "Table Grid"
    for j, h in enumerate(["Instrumento normativo", "Relevancia para el proyecto"]):
        c = tbl_n.rows[0].cells[j]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (norma, relevancia) in enumerate(normas, 1):
        tbl_n.rows[i].cells[0].text = norma
        tbl_n.rows[i].cells[1].text = relevancia
        for j in range(2):
            tbl_n.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_page_break()

    # ━━━ 4. DEFINICIONES Y ABREVIATURAS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "4. DEFINICIONES Y ABREVIATURAS", 1)

    _h(doc, "4.1 Definiciones", 2)
    definiciones = [
        ("Kiosko comercial",
         "Módulo modular de venta de productos o servicios, de dimensiones reducidas "
         "(2.5 m × 2.0 m), diseñado para instalación en espacios públicos con mínimo "
         "impacto paisajístico. Puede ser fijo, semi-permanente o móvil."),
        ("Bulevar de las Canchas",
         "Corredor peatonal de 160 metros lineales que atraviesa longitudinalmente el "
         "área de canchas deportivas del Parque Bicentenario, identificado como la zona "
         "de mayor flujo peatonal y la más adecuada para la implantación de kioskos."),
        ("Giro comercial",
         "Categoría de productos o servicios que ofrece un kiosko. Ejemplos: Bebidas, "
         "Comida rápida, Helados y café, Snacks saludables, Artículos para mascotas y "
         "souvenirs, Deportivo."),
        ("Rotación de giros",
         "Sistema de asignación de categorías comerciales a cada kiosko que garantiza "
         "diversidad de oferta y evita competencia directa entre módulos adyacentes."),
        ("AIVA (Área de Intervención Valorativa)",
         "Indicador del valor del suelo urbano expresado en dólares por metro cuadrado "
         "(USD/m²). En el sector Bicentenario el AIVA actual es de $267/m², mientras que "
         "en el sector de referencia La Carolina asciende a $1,895/m²."),
        ("Demanda potencial",
         "Número de visitantes que declararon que consumirían productos o servicios "
         "comerciales dentro del parque si estos estuvieran disponibles."),
        ("Visitante activo",
         "Persona que ingresa al parque para realizar actividad física, deportiva o "
         "recreativa, con mayor propensión al consumo de bebidas e hidratantes."),
        ("Plan de implementación por fases",
         "Estrategia gradual de despliegue de kioskos en el tiempo, condicionada a "
         "hitos de desarrollo urbano del sector (apertura del Metro, densificación "
         "de edificaciones, crecimiento poblacional)."),
    ]
    tbl_d = doc.add_table(rows=len(definiciones) + 1, cols=2)
    tbl_d.style = "Table Grid"
    for j, h in enumerate(["Término", "Definición"]):
        c = tbl_d.rows[0].cells[j]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (term, defi) in enumerate(definiciones, 1):
        tbl_d.rows[i].cells[0].text = term
        tbl_d.rows[i].cells[1].text = defi
        tbl_d.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl_d.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        tbl_d.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    _h(doc, "4.2 Abreviaturas", 2)
    abreviaturas = [
        ("AIVA",    "Área de Intervención Valorativa"),
        ("COA",     "Código Orgánico del Ambiente"),
        ("COOTAD",  "Código Orgánico de Organización Territorial, Autonomía y Descentralización"),
        ("DMQ",     "Distrito Metropolitano de Quito"),
        ("EPMMOP",  "Empresa Pública Metropolitana de Movilidad y Obras Públicas"),
        ("KPI",     "Key Performance Indicator — Indicador Clave de Rendimiento"),
        ("MDMQ",    "Municipio del Distrito Metropolitano de Quito"),
        ("PUGS",    "Plan de Uso y Gestión del Suelo"),
        ("USD",     "Dólares de los Estados Unidos de América"),
    ]
    tbl_a = doc.add_table(rows=len(abreviaturas) + 1, cols=2)
    tbl_a.style = "Table Grid"
    for j, h in enumerate(["Abreviatura", "Significado"]):
        c = tbl_a.rows[0].cells[j]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (abr, sig) in enumerate(abreviaturas, 1):
        tbl_a.rows[i].cells[0].text = abr
        tbl_a.rows[i].cells[1].text = sig
        tbl_a.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl_a.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        tbl_a.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_page_break()

    # ━━━ 5. DESARROLLO ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "5. DESARROLLO", 1)

    # 5.1 Metodología
    _h(doc, "5.1 Metodología", 2)
    _p(doc, (
        "El estudio combina métodos cuantitativos y cualitativos en dos etapas: "
        "(a) encuesta directa a visitantes del parque y (b) análisis estadístico de "
        "los registros históricos de visita proporcionados por la EPMMOP."
    ))
    _kpi_table(doc, [
        ("Tamaño de muestra",        f"{n_enc} encuestados"),
        ("Herramienta de recolección","Google Forms"),
        ("Tipo de muestreo",         "Aleatorio sistemático"),
        ("Variables analizadas",     "17 preguntas"),
        ("Registros de visita",      f"{vis_2025:,} (año 2025)"),
        ("Herramienta de análisis",  "Python — pandas, scipy, matplotlib"),
    ])

    # 5.2 Perfil sociodemográfico
    _h(doc, "5.2 Perfil sociodemográfico del visitante", 2)
    _p(doc, (
        f"El visitante típico del Parque Bicentenario es residente del norte de Quito (78.2%), "
        f"con una edad promedio de {edad} años. El domingo es el día de mayor afluencia (67.5%), "
        "con horario pico de 09:00 a 12:00 h en fines de semana y de 04:00 a 09:00 h en días "
        "laborables. El grupo de visita predominante es la familia (45.6%) y el motivo "
        "principal es la recreación (29.1%)."
    ))
    _kpi_table(doc, [
        ("Visitantes totales 2025",            f"{vis_2025:,}"),
        ("Visitantes diarios promedio",        f"{vis_dia:,}"),
        ("Edad promedio",                      f"{edad} años"),
        ("Día de mayor afluencia",             "Domingo (67.5%)"),
        ("Sector de residencia predominante",  "Norte de Quito (78.2%)"),
        ("Grupo de visita predominante",       "Familia (45.6%)"),
        ("Motivo de visita principal",         "Recreación (29.1%)"),
        ("Servicio complementario prioritario","Baños públicos (76.2%)"),
    ])
    _insert_img(doc, charts.get("genero"),
                "Gráfico 1. Distribución por género de los visitantes encuestados", 4.5)
    _insert_img(doc, charts.get("edad"),
                "Gráfico 2. Distribución por grupos de edad", 5.0)
    _insert_img(doc, charts.get("motivo"),
                "Gráfico 3. Motivo principal de visita al Parque Bicentenario", 4.5)

    # 5.3 Análisis de la demanda
    _h(doc, "5.3 Análisis de la demanda", 2)
    _p(doc, (
        f"El Parque Bicentenario registró {vis_2025:,} visitas durante 2025, equivalentes a un "
        f"promedio de {vis_dia:,} visitantes diarios. La tendencia mensual muestra estacionalidad "
        "en julio–agosto (vacaciones escolares) y en el primer trimestre del año."
    ))
    _insert_img(doc, charts.get("trafico"),
                "Gráfico 4. Tráfico mensual de visitantes al Parque Bicentenario — 2025", 6.0)

    _p(doc, (
        f"Del total de visitantes, el {cons_pct}% ({cons_dia} personas/día) declaró que "
        f"consumiría productos o servicios si estuvieran disponibles dentro del parque. "
        f"El gasto promedio dispuesto es de ${gasto} por visita. El {aprueba}% aprueba la "
        f"implementación de kioskos y el {mejora}% considera que mejoraría su experiencia."
    ))
    _insert_img(doc, charts.get("demanda"),
                "Gráfico 5. Demanda potencial de consumidores vs. capacidad del plan de kioskos", 6.0)

    _p(doc, (
        "Los productos y servicios que los visitantes preferirían adquirir dentro del parque "
        "se identifican a partir de la pregunta directa de la encuesta. El gráfico siguiente "
        "muestra la distribución porcentual de preferencias por categoría de producto:"
    ))
    _insert_img(doc, charts.get("productos"),
                "Gráfico 6. Productos y servicios más demandados por los visitantes del parque", 6.0)

    _p(doc, (
        "Complementariamente, el análisis de los comentarios abiertos de los visitantes permite "
        "identificar las necesidades y mejoras más mencionadas de forma espontánea. Este análisis "
        "revela las prioridades percibidas por la ciudadanía más allá de las opciones cerradas "
        "de la encuesta:"
    ))
    _insert_img(doc, charts.get("topics"),
                "Gráfico 7. Temas más mencionados en comentarios de visitantes (% de menciones)", 6.0)

    # 5.4 Propuesta de modelo comercial
    _h(doc, "5.4 Propuesta de modelo comercial", 2)
    _p(doc, (
        "La zona de intervención seleccionada es el Bulevar de las Canchas (Zona 1), "
        "corredor peatonal de 160 m de longitud que atraviesa el área de canchas deportivas "
        "del parque. La elección de esta zona obedece a su alta concentración de flujo "
        "peatonal, su conexión directa con los accesos norte y sur, y su compatibilidad "
        "con el uso recreativo y deportivo del sector."
    ))
    _kpi_table(doc, [
        ("Zona de intervención",  "Zona 1 – Bulevar de las Canchas"),
        ("Longitud total",        "160 m lineales — 4 secciones"),
        ("Capacidad máxima",      "50 kioskos"),
        ("Dimensión por kiosko",  "2.5 m × 2.0 m"),
        ("Ancho del pasillo",     "6.0 m"),
        ("Tipo de módulo",        "Semi-permanente / desmontable"),
    ])

    zona_img = "src/img/render_bicentenario_bulevar.png"
    if os.path.exists(zona_img):
        doc.add_picture(zona_img, width=Inches(6.0))
        cap = doc.add_paragraph("Figura 1. Zona de intervención — Bulevar de las Canchas")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].italic = True
        doc.add_paragraph()

    prop_img = "src/img/implantación_parque.jpg"
    if os.path.exists(prop_img):
        doc.add_picture(prop_img, width=Inches(6.0))
        cap2 = doc.add_paragraph("Figura 2. Render de la propuesta — Parque Bicentenario")
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap2.runs[0].font.size = Pt(9)
        cap2.runs[0].italic = True
        doc.add_paragraph()

    _p(doc, "Giros comerciales propuestos y distribución por rotación:", bold=True)
    giro_rows = [(g, GIROS[g]["descripcion"]) for g in GIROS]
    tbl_g = doc.add_table(rows=len(giro_rows) + 1, cols=2)
    tbl_g.style = "Table Grid"
    for j, h in enumerate(["Giro comercial", "Productos / servicios"]):
        c = tbl_g.rows[0].cells[j]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (giro, desc) in enumerate(giro_rows, 1):
        tbl_g.rows[i].cells[0].text = giro
        tbl_g.rows[i].cells[1].text = desc
        for j in range(2):
            tbl_g.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    _p(doc, "Rotación de giros — Fase 1 (10 kioskos):", bold=True)
    rotation = GIRO_ROTATION["comercial_alta_densidad"]
    tbl_r = doc.add_table(rows=2, cols=10)
    tbl_r.style = "Table Grid"
    for i, giro in enumerate(rotation[:10]):
        tbl_r.rows[0].cells[i].text = f"K{i+1}"
        tbl_r.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        tbl_r.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(8)
        tbl_r.rows[1].cells[i].text = giro
        tbl_r.rows[1].cells[i].paragraphs[0].runs[0].font.size = Pt(8)
    doc.add_paragraph()

    _insert_img(doc, charts.get("layout"),
                "Figura 3. Distribución comercial propuesta — Bulevar de las Canchas, Fase 1", 6.0)

    # Plan de fases
    _p(doc, "Plan de implementación por fases:", bold=True)
    fase_rows = []
    for anio, params in sorted(PLAN_FASES.items()):
        total_kio = sum(params["kioskos"].values())
        fase_rows.append((params["fase"], str(anio), params["hito"],
                          str(total_kio), f"{params['poblacion_hab']:,} hab."))
    tbl_f = doc.add_table(rows=len(fase_rows) + 1, cols=5)
    tbl_f.style = "Table Grid"
    for j, h in enumerate(["Fase", "Año", "Hito clave", "Kioskos", "Población sector"]):
        c = tbl_f.rows[0].cells[j]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
    for i, row in enumerate(fase_rows, 1):
        for j, val in enumerate(row):
            tbl_f.rows[i].cells[j].text = val
            tbl_f.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    _insert_img(doc, charts.get("kioskos"),
                "Gráfico 8. Proyección de kioskos por fase — 2026 a 2036", 6.0)

    # 5.5 Proyección financiera
    _h(doc, "5.5 Proyección financiera y viabilidad", 2)
    _p(doc, (
        f"La proyección se basa en el crecimiento esperado de la población del sector "
        f"({pob_2026:,} hab. en 2026 → {pob_2036:,} hab. en 2036), el aumento progresivo "
        "de visitantes por la apertura del Metro (2029) y el incremento gradual de kioskos "
        "activos. Los modelos de ingreso utilizan el ticket promedio declarado por los "
        "encuestados (${gasto}) bajo escenarios conservadores de ocupación.".format(gasto=gasto)
    ))
    _insert_img(doc, charts.get("visitantes_vs"),
                "Gráfico 9. Proyección de visitantes vs. consumidores potenciales por fase", 6.0)
    _insert_img(doc, charts.get("ing_zona"),
                "Gráfico 10. Proyección de ingresos totales de la zona 2026–2036 (miles USD/año)", 6.0)
    _insert_img(doc, charts.get("ing_kio"),
                "Gráfico 11. Proyección de ingresos promedio por kiosko 2026–2036 (miles USD/año)", 6.0)

    _p(doc, "Impacto en valorización del suelo (AIVA):", bold=True)
    _kpi_table(doc, [
        ("AIVA actual sector Bicentenario",  "$267/m²"),
        ("AIVA referencia sector La Carolina","$1,895/m²"),
        ("Potencial de crecimiento",         "~7x respecto al valor actual"),
        ("Beneficio para propietarios",      "Incremento de valor de venta y arrendamiento "
                                             "de inmuebles conforme el proyecto avanza"),
    ])
    doc.add_page_break()

    # ━━━ 6. CONCLUSIONES ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "6. CONCLUSIONES", 1)
    _numbered(doc, [
        f"Existe demanda potencial verificada: el {cons_pct}% de los visitantes ({cons_dia} "
        f"personas/día) consumiría productos o servicios si estuvieran disponibles en el parque.",
        f"El parque recibe {vis_2025:,} visitas anuales ({vis_dia:,}/día promedio), con mayor "
        f"concentración los domingos. La apertura del Metro (2029) incrementará significativamente "
        "esta afluencia.",
        "Los snacks, bebidas y comida rápida son los productos más demandados (38.0%); "
        "los baños públicos son la prioridad de servicio complementario más urgente (76.2%).",
        f"El modelo de kioskos por fases (10 → 25 unidades, 2026–2036) es técnicamente viable, "
        "escalable y coherente con el crecimiento urbano proyectado para el sector.",
        f"La ciudadanía tiene una actitud favorable hacia el proyecto: {aprueba}% aprueba los "
        f"kioskos y {mejora}% considera que mejorarían su experiencia en el parque.",
        "El AIVA del sector Bicentenario ($267/m²) presenta un alto potencial de valorización "
        "hacia los niveles del sector La Carolina ($1,895/m²), generando un beneficio colateral "
        "para los propietarios de inmuebles del área de influencia.",
        "El Bulevar de las Canchas, con 160 m lineales y capacidad para hasta 50 módulos, "
        "es la zona técnicamente idónea para iniciar la implementación del proyecto.",
    ])
    doc.add_page_break()

    # ━━━ 7. RECOMENDACIONES ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "7. RECOMENDACIONES", 1)
    _numbered(doc, [
        "Iniciar la Fase 1 del proyecto en 2026 con 10 kioskos de giros diversificados "
        "(Bebidas, Comida rápida, Helados y café, Snacks saludables, Mascotas/souvenirs y "
        "Deportivo), para validar el modelo comercial antes de escalar.",
        "Priorizar la instalación de baños públicos como servicio complementario urgente, "
        "en respuesta al 76.2% de visitantes que los identifica como necesidad principal.",
        "Establecer contratos de concesión con cláusulas de revisión anual, que permitan "
        "ajustar los giros comerciales y las tarifas en función de la demanda real observada.",
        "Implementar un sistema de monitoreo de ventas y flujo de visitantes que permita "
        "tomar decisiones de escalamiento basadas en datos en tiempo real.",
        "Estructurar una campaña de comunicación hacia propietarios e inversores del sector "
        "sobre el impacto positivo del proyecto en el AIVA, para generar alianzas ciudadanas.",
        "Programar la transición a la Fase 2 (16 kioskos) en coordinación con la apertura "
        "de las estaciones de Metro Bicentenario y Andalucía (2029).",
        "Evaluar la extensión del modelo a otras zonas del parque en la Fase 4 (2036) "
        "si los resultados de las fases previas superan las proyecciones conservadoras.",
    ])
    doc.add_page_break()

    # ━━━ 8. FIRMAS DE RESPONSABILIDAD ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "8. FIRMAS DE RESPONSABILIDAD", 1)
    _p(doc, (
        "El presente estudio de mercado fue elaborado, revisado y aprobado por los "
        "funcionarios de la EPMMOP que a continuación suscriben, quienes certifican la "
        "veracidad y rigurosidad técnica de la información contenida en este documento."
    ))
    doc.add_paragraph()

    _firma_table(doc, [
        ("Elaborado por", ""),
        ("Revisado por",  ""),
        ("Aprobado por",  ""),
    ])

    doc.add_paragraph()
    _kpi_table(doc, [
        ("Elaborado por",  "Gerencia de Promoción de Servicios — EPMMOP"),
        ("Revisado por",   "Gerencia de Desarrollo Urbano — EPMMOP"),
        ("Aprobado por",   "Gerencia General — EPMMOP"),
        ("Fecha",          date.today().strftime("%d de %B de %Y").capitalize()),
        ("N.° de documento","EPMMOP-PS-EM-2026-001"),
    ])
    doc.add_page_break()

    # ━━━ 9. ANEXOS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    _h(doc, "9. ANEXOS", 1)

    _h(doc, "Anexo A — Indicadores clave del estudio (KPIs)", 2)
    _kpi_table(doc, [
        ("Visitas registradas 2025",           f"{vis_2025:,}"),
        ("Visitantes diarios promedio",        f"{vis_dia:,}"),
        ("Encuestados",                        f"{n_enc}"),
        ("Consumidores potenciales por día",   f"{cons_dia} ({cons_pct}%)"),
        ("Gasto promedio dispuesto",           f"${gasto}"),
        ("Aprobación de kioskos",              f"{aprueba}%"),
        ("Mejora percibida en experiencia",    f"{mejora}%"),
        ("Calificación actual del parque",     f"{calif}/5"),
        ("Edad promedio del visitante",        f"{edad} años"),
        ("AIVA sector Bicentenario",           "$267/m²"),
        ("AIVA referencia La Carolina",        "$1,895/m²"),
        ("Kioskos inicio / meta",              "10 (2026) → 25 (2036)"),
    ])

    _h(doc, "Anexo B — Instrumento de encuesta", 2)
    _p(doc, f"Variables del formulario aplicado a {n_enc} visitantes del Parque Bicentenario:")
    _numbered(doc, [
        "Edad",
        "Género",
        "¿En qué sector reside?",
        "¿Con quién visita generalmente el Parque Bicentenario?",
        "¿Con qué frecuencia visita el Parque Bicentenario?",
        "¿Qué días usualmente visita el parque?",
        "¿En qué horario usualmente visita el parque?",
        "¿Cuál es el principal motivo de su visita?",
        "¿Consumiría productos o servicios dentro del parque?",
        "¿Qué productos o servicios consumiría dentro del parque?",
        "¿Cuánto estaría dispuesto a gastar en la compra de estos productos?",
        "¿Considera adecuada la implementación de kioskos comerciales dentro del parque?",
        "¿En qué zonas accedería con mayor facilidad para adquirir productos?",
        "¿Considera que los kioskos mejorarían su experiencia en el parque?",
        "¿Cómo califica actualmente la oferta de servicios dentro del parque?",
        "¿Qué servicios complementarios considera prioritarios?",
        "Comentarios y sugerencias",
    ])

    _h(doc, "Anexo C — Descripción detallada del plan de implementación por fases", 2)
    for anio, params in sorted(PLAN_FASES.items()):
        _p(doc, f"{params['fase']} — {anio}", bold=True, size=10)
        _p(doc, f"Hito: {params['hito']}", size=10)
        _p(doc, params["contexto"], size=10)
        _p(doc, f"Kioskos: {sum(params['kioskos'].values())} | "
                f"Población sector: {params['poblacion_hab']:,} hab.", size=10)
        doc.add_paragraph()

    # ── guardar ───────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
